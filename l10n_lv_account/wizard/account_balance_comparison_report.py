# -*- encoding: utf-8 -*-
##############################################################################
#
#    OpenERP, Open Source Management Solution
#    Copyright (C) 2016 ITS-1 (<http://www.its1.lv/>)
#                       E-mail: <info@its1.lv>
#                       Address: <Vienibas gatve 109 LV-1058 Riga Latvia>
#                       Phone: +371 66116534
#
#    This program is free software: you can redistribute it and/or modify
#    it under the terms of the GNU Affero General Public License as
#    published by the Free Software Foundation, either version 3 of the
#    License, or (at your option) any later version.
#
#    This program is distributed in the hope that it will be useful,
#    but WITHOUT ANY WARRANTY; without even the implied warranty of
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#    GNU Affero General Public License for more details.
#
#    You should have received a copy of the GNU Affero General Public License
#    along with this program.  If not, see <http://www.gnu.org/licenses/>.
#
##############################################################################

from openerp.osv import fields, osv
from openerp.tools.translate import _
import time
import base64
import StringIO
import xlwt
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from openerp.addons.l10n_lv_account.report.report_account_balance_comparison import report_account_balance_comparison as rabc
from openerp.report.report_sxw import rml_parse

class account_balance_comparison_report(osv.osv_memory):
    _name = "account.balance.comparison.report"
    _description = "Partner Balance Comparison"

    _columns = {
        'date': fields.date('Date', required=True),
        'type': fields.selection([('receivable', 'Receivable'), ('payable', 'Payable')], string='Type', required=True),
        'accountant_id': fields.many2one('hr.employee', 'Accountant', required=True),
        'format': fields.selection([('pdf','PDF'), ('docx','DOCX')], string='Format', required=True),
        'file_name': fields.char('File Name'),
        'file_save': fields.binary('Save File', readonly=True)
    }

    def _get_default_accountant(self, cr, uid, context=None):
        emp_obj = self.pool.get('hr.employee')
        emp_ids = emp_obj.search(cr, uid, [('job_id.name','in',['Accountant', 'accountant', 'Bookkeeper', 'bookkeeper'])], context=context)
        if not emp_ids:
            emp_ids = emp_obj.search(cr, uid, ['|', ('job_id.name','ilike','accountant'), ('job_id.name','ilike','bookkeeper')], context=context)
        return emp_ids and emp_ids[0] or False

    def _get_default_type(self, cr, uid, context=None):
        if context is None:
            context = {}
        type = False
        if context.get('search_default_customer',False):
            type = 'receivable'
        if context.get('search_default_supplier',False):
            type = 'payable'
        if not type:
            types = []
            p_ids = context.get('active_ids',[])
            for p in self.pool.get('res.partner').browse(cr, uid, p_ids, context=context):
                if p.customer == True and 'receivable' not in types:
                    types.append('receivable')
                if p.supplier == True and 'payable' not in types:
                    types.append('payable')
            if len(types) == 1:
                type = types[0]
        return type

    _defaults = {
        'date': fields.date.context_today,
        'type': _get_default_type,
        'accountant_id': _get_default_accountant,
        'format': 'pdf',
        'file_name': 'Payment_Comparison.docx'
    }

    def _build_contexts(self, cr, uid, ids, data, context=None):
        if context is None:
            context = {}
        result = {}
        result['date'] = 'date' in data['form'] and data['form']['date'] or False
        result['type'] = 'date' in data['form'] and data['form']['type'] or False
        result['accountant_id'] = 'accountant_id' in data['form'] and data['form']['accountant_id'] or False
        result['format'] = 'format' in data['form'] and data['form']['format'] or False
        result['file_name'] = 'format' in data['form'] and data['form']['file_name'] or False
        return result

    def get_transl(self, cr, uid, string, lang, context=None):
        if context is None:
            context = {}
        lang = lang or context.get('lang',False)
        data_obj = self.pool.get('ir.model.data')
        report = data_obj.get_object(cr, uid, 'l10n_lv_account', 'balance_comparison_document', context=context)
        transl_obj = self.pool.get('ir.translation')
        transl_ids = transl_obj.search(cr, uid, [('name','=','website'), ('module','=','l10n_lv_account'), ('type','=','view'), ('res_id','=',report.id), ('src','=',string), ('lang','=',lang)], context=context)
        if transl_ids:
            transl = transl_obj.browse(cr, uid, transl_ids[0], context=context)
            return transl.value
        return string

    def make_docx_data(self, cr, uid, partner_ids, data, context=None):
        if context is None:
            context = {}
        user = self.pool.get('res.users').browse(cr, uid, uid, context=context)
        user_company = user.company_id
        partner_obj = self.pool.get('res.partner')
        document = docx.Document()
        for partner in partner_obj.browse(cr, uid, partner_ids, context=context):
            company = partner.company_id and partner.company_id or user_company
            section = document.sections[-1]
            section.start_type = WD_SECTION.ODD_PAGE
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            h2_str = self.get_transl(cr, uid, "Mutual payment comparison statement", partner.lang, context=context)
            h2 = document.add_paragraph()
            h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h2_run = h2.add_run(h2_str)
            h2_run.bold = True
            h2_run.font.name = 'Liberation Sans'
            h2_run.font.size = Pt(14)
            space1 = document.add_paragraph()
            space1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            space1_run = space1.add_run()
            space1_run.bold = True
            space1_run.font.name = 'Liberation Sans'
            space1_run.font.size = Pt(14)
#            space2 = document.add_paragraph()
#            space2_run = space1.add_run()
#            space2_run.font.name = 'Liberation Sans'
#            space2_run.font.size = Pt(12)

            # Info Table:
            info_table = document.add_table(rows=6, cols=3)
            info_table.style = 'TableGrid'
            info_table.columns[0].width = Inches(3.67)
            info_table.columns[1].width = Inches(0.07)
            info_table.columns[2].width = Inches(3.74)

            # row0 column0:
            itr0c0 = info_table.rows[0].cells[0].add_paragraph()
            itr0c0_run1_str = self.get_transl(cr, uid, "Company:", partner.lang, context=context) + ' '
            itr0c0_run1 = itr0c0.add_run(itr0c0_run1_str)
            itr0c0_run1.font.name = 'Liberation Sans'
            itr0c0_run1.font.size = Pt(12)
            itr0c0_run2 = itr0c0.add_run(company.name)
            itr0c0_run2.bold = True
            itr0c0_run2.font.name = 'Liberation Sans'
            itr0c0_run2.font.size = Pt(12)

            # row0 column1:
            info_table.rows[0].cells[1].merge(info_table.rows[1].cells[1]).merge(info_table.rows[2].cells[1]).merge(info_table.rows[3].cells[1]).merge(info_table.rows[4].cells[1])

            # row0 column2:
            itr0c2 = info_table.rows[0].cells[2].add_paragraph()
            itr0c2_run1 = itr0c2.add_run(itr0c0_run1_str)
            itr0c2_run1.font.name = 'Liberation Sans'
            itr0c2_run1.font.size = Pt(12)
            itr0c2_run2 = itr0c2.add_run(partner.name)
            itr0c2_run2.bold = True
            itr0c2_run2.font.name = 'Liberation Sans'
            itr0c2_run2.font.size = Pt(12)

            # row1 column0:
            itr1c0_1 = info_table.rows[1].cells[0].add_paragraph()
            itr1c0_1_run_str = rabc(cr, uid, '', context).form_address(company)
            itr1c0_1_run = itr1c0_1.add_run(itr1c0_1_run_str)
            itr1c0_1_run.font.name = 'Tahoma'
            itr1c0_1_run.font.size = Pt(10)
            if company.phone:
                itr1c0_2 = info_table.rows[1].cells[0].add_paragraph()
                itr1c0_2_run_str = self.get_transl(cr, uid, "Phone:", partner.lang, context=context) + ' ' + company.phone
                itr1c0_2_run = itr1c0_2.add_run(itr1c0_2_run_str)
                itr1c0_2_run.font.name = 'Tahoma'
                itr1c0_2_run.font.size = Pt(10)
            if company.email:
                itr1c0_3 = info_table.rows[1].cells[0].add_paragraph()
                itr1c0_3_run_str = self.get_transl(cr, uid, "E-mail:", partner.lang, context=context) + ' ' + company.email
                itr1c0_3_run = itr1c0_3.add_run(itr1c0_3_run_str)
                itr1c0_3_run.font.name = 'Tahoma'
                itr1c0_3_run.font.size = Pt(10)

            # row1 column2:
            itr1c2_1 = info_table.rows[1].cells[2].add_paragraph()
            itr1c2_1_run_str = rabc(cr, uid, '', context).form_address(partner)
            itr1c2_1_run = itr1c2_1.add_run(itr1c2_1_run_str)
            itr1c2_1_run.font.name = 'Tahoma'
            itr1c2_1_run.font.size = Pt(10)
            if partner.phone:
                itr1c2_2 = info_table.rows[1].cells[2].add_paragraph()
                itr1c2_2_run_str = self.get_transl(cr, uid, "Phone:", partner.lang, context=context) + ' ' + partner.phone
                itr1c2_2_run = itr1c2_2.add_run(itr1c2_2_run_str)
                itr1c2_2_run.font.name = 'Tahoma'
                itr1c2_2_run.font.size = Pt(10)
            if partner.email:
                itr1c2_3 = info_table.rows[1].cells[2].add_paragraph()
                itr1c2_3_run_str = self.get_transl(cr, uid, "E-mail:", partner.lang, context=context) + ' ' + partner.email
                itr1c2_3_run = itr1c2_3.add_run(itr1c2_3_run_str)
                itr1c2_3_run.font.name = 'Tahoma'
                itr1c2_3_run.font.size = Pt(10)

            # row2 column0:
            itr2c0 = info_table.rows[2].cells[0].add_paragraph()
            itr2c0_run_str = self.get_transl(cr, uid, "TIN:", partner.lang, context=context) + ' ' + company.vat
            itr2c0_run = itr2c0.add_run(itr2c0_run_str)
            itr2c0_run.font.name = 'Liberation Sans'
            itr2c0_run.font.size = Pt(10)

            # row2 column2:
            itr2c2 = info_table.rows[2].cells[2].add_paragraph()
            itr2c2_run_str = self.get_transl(cr, uid, "TIN:", partner.lang, context=context) + ' ' + partner.vat
            itr2c2_run = itr2c2.add_run(itr2c2_run_str)
            itr2c2_run.font.name = 'Liberation Sans'
            itr2c2_run.font.size = Pt(10)

            # row3 column0:
            itr3c0_1 = info_table.rows[3].cells[0].add_paragraph()
            itr3c0_1_run_str = self.get_transl(cr, uid, "Accountant:", partner.lang, context=context) + ' ' + rabc(cr, uid, '', context).get_accountant(data['form']['accountant_id'])
            itr3c0_1_run = itr3c0_1.add_run(itr3c0_1_run_str)
            itr3c0_1_run.font.name = 'Liberation Sans'
            itr3c0_1_run.font.size = Pt(10)

            # row3 column2:
            itr3c2_1 = info_table.rows[3].cells[2].add_paragraph()
            itr3c2_1_run_str = self.get_transl(cr, uid, "Accountant:", partner.lang, context=context)
            itr3c2_1_run = itr3c2_1.add_run(itr3c2_1_run_str)
            itr3c2_1_run.font.name = 'Liberation Sans'
            itr3c2_1_run.font.size = Pt(10)
            itr3c2_2 = info_table.rows[3].cells[2].add_paragraph()
            itr3c2_2_run = itr3c2_2.add_run()
            itr3c2_2_run.font.name = 'Liberation Sans'
            itr3c2_2_run.font.size = Pt(10)
            itr3c2_3 = info_table.rows[3].cells[2].add_paragraph()
            itr3c2_3_run_str = "_______________________________________"
            itr3c2_3_run = itr3c2_3.add_run(itr3c2_3_run_str)
            itr3c2_3_run.font.name = 'Liberation Sans'
            itr3c2_3_run.font.size = Pt(10)
            itr3c2_3_run.font.color.rgb = RGBColor(0xD9, 0xD9, 0xD9)

            # row4 column0:
            itr4c0_1 = info_table.rows[4].cells[0].add_paragraph()
            itr4c0_1_run_str = self.get_transl(cr, uid, "Date:", partner.lang, context=context)
            itr4c0_1_run = itr4c0_1.add_run(itr4c0_1_run_str)
            itr4c0_1_run.bold = True
            itr4c0_1_run.font.name = 'Liberation Sans'
            itr4c0_1_run.font.size = Pt(12)
            itr4c0_2 = info_table.rows[4].cells[0].add_paragraph()
            itr4c0_2_run_str = rml_parse(cr, uid, '', context).formatLang(time.strftime('%Y-%m-%d'), date=True)
            itr4c0_2_run = itr4c0_2.add_run(itr4c0_2_run_str)
            itr4c0_2_run.font.name = 'Liberation Sans'
            itr4c0_2_run.font.size = Pt(12)

            # row4 column2:
            itr4c2_1 = info_table.rows[4].cells[2].add_paragraph()
            itr4c2_1_run_str = self.get_transl(cr, uid, "Date:", partner.lang, context=context)
            itr4c2_1_run = itr4c2_1.add_run(itr4c2_1_run_str)
            itr4c2_1_run.bold = True
            itr4c2_1_run.font.name = 'Liberation Sans'
            itr4c2_1_run.font.size = Pt(12)
            itr4c2_2 = info_table.rows[4].cells[2].add_paragraph()
            itr4c2_2_run1 = itr4c2_2.add_run('20')
            itr4c2_2_run1.font.name = 'Liberation Sans'
            itr4c2_2_run1.font.size = Pt(12)
            itr4c2_2_run2 = itr4c2_2.add_run('______')
            itr4c2_2_run2.font.name = 'Liberation Sans'
            itr4c2_2_run2.font.size = Pt(12)
            itr4c2_2_run2.font.color.rgb = RGBColor(0xD9, 0xD9, 0xD9)
            itr4c2_2_run3 = itr4c2_2.add_run(' "')
            itr4c2_2_run3.font.name = 'Liberation Sans'
            itr4c2_2_run3.font.size = Pt(12)
            itr4c2_2_run4 = itr4c2_2.add_run('_______')
            itr4c2_2_run4.font.color.rgb = RGBColor(0xD9, 0xD9, 0xD9)
            itr4c2_2_run4.font.name = 'Liberation Sans'
            itr4c2_2_run4.font.size = Pt(12)
            itr4c2_2_run5 = itr4c2_2.add_run('".')
            itr4c2_2_run5.font.name = 'Liberation Sans'
            itr4c2_2_run5.font.size = Pt(12)
            itr4c2_2_run6 = itr4c2_2.add_run('_____________________\n')
            itr4c2_2_run6.font.name = 'Liberation Sans'
            itr4c2_2_run6.font.size = Pt(12)
            itr4c2_2_run6.font.color.rgb = RGBColor(0xD9, 0xD9, 0xD9)

            # row5 column0:
            itr5c0 = info_table.rows[5].cells[0].add_paragraph()
            itr5c0_run1_str = self.get_transl(cr, uid, "The turnover check", partner.lang, context=context) + ' '
            itr5c0_run1 = itr5c0.add_run(itr5c0_run1_str)
            itr5c0_run1.font.name = 'Liberation Sans'
            itr5c0_run1.font.size = Pt(12)
            itr5c0_run2_str = self.get_transl(cr, uid, "recognised", partner.lang, context=context)
            itr5c0_run2 = itr5c0.add_run(itr5c0_run2_str)
            itr5c0_run2.bold = True
            itr5c0_run2.font.name = 'Liberation Sans'
            itr5c0_run2.font.size = Pt(12)
            itr5c0_run3 = itr5c0.add_run(':')
            itr5c0_run3.font.name = 'Liberation Sans'
            itr5c0_run3.font.size = Pt(12)
            info_table.rows[5].cells[0].merge(info_table.rows[5].cells[1]).merge(info_table.rows[5].cells[2])

            space3 = document.add_paragraph()
            space3_run = space3.add_run()
            space3_run.font.name = 'Liberation Sans'
            space3_run.font.size = Pt(12)

            # List Table:
            list_table = document.add_table(rows=7, cols=5)
            list_table.style = 'TableGrid'
            list_table.columns[0].width = Inches(1.3)
            list_table.columns[1].width = Inches(1.23)
            list_table.columns[2].width = Inches(2.41)
            list_table.columns[3].width = Inches(1.24)
            list_table.columns[4].width = Inches(1.3)

            # row0 column0:
            ltr0c0 = list_table.rows[0].cells[0].add_paragraph()
            ltr0c0_run_str = self.get_transl(cr, uid, "According to:", partner.lang, context=context)
            ltr0c0_run = ltr0c0.add_run(ltr0c0_run_str)
            ltr0c0_run.bold = True
            ltr0c0_run.font.name = 'Liberation Sans'
            ltr0c0_run.font.size = Pt(10)
            list_table.rows[0].cells[0].merge(list_table.rows[0].cells[1])

            # row0 column2:
            ltr0c2 = list_table.rows[0].cells[2].add_paragraph()
            ltr0c2_run_str = self.get_transl(cr, uid, "Balance on:", partner.lang, context=context)
            ltr0c2_run = ltr0c2.add_run(ltr0c2_run_str)
            ltr0c2_run.bold = True
            ltr0c2_run.font.name = 'Liberation Sans'
            ltr0c2_run.font.size = Pt(10)

            # row0 column3:
            ltr0c3 = list_table.rows[0].cells[3].add_paragraph()
            ltr0c3_run_str = self.get_transl(cr, uid, "According to:", partner.lang, context=context)
            ltr0c3_run = ltr0c3.add_run(ltr0c3_run_str)
            ltr0c3_run.bold = True
            ltr0c3_run.font.name = 'Liberation Sans'
            ltr0c3_run.font.size = Pt(10)
            list_table.rows[0].cells[3].merge(list_table.rows[0].cells[4])

            # row1 column0:
            ltr1c0 = list_table.rows[1].cells[0].add_paragraph()
            ltr1c0_run = ltr1c0.add_run(company.name)
            ltr1c0_run.bold = True
            ltr1c0_run.font.name = 'Liberation Sans'
            ltr1c0_run.font.size = Pt(10)
            list_table.rows[1].cells[0].merge(list_table.rows[1].cells[1])

            # row1 column2:
            ltr1c2_1 = list_table.rows[1].cells[2].add_paragraph()
            ltr1c2_1_run = ltr1c2_1.add_run()
            ltr1c2_1_run.font.name = 'Liberation Sans'
            ltr1c2_1_run.font.size = Pt(10)
            ltr1c2_2 = list_table.rows[1].cells[2].add_paragraph()
            ltr1c2_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ltr1c2_2_run_str = rml_parse(cr, uid, '', context).formatLang(time.strftime('%Y-%m-%d'), date=True)
            ltr1c2_2_run = ltr1c2_2.add_run(ltr1c2_2_run_str)
            ltr1c2_2_run.font.name = 'Liberation Sans'
            ltr1c2_2_run.font.size = Pt(11)
            list_table.rows[1].cells[2].merge(list_table.rows[2].cells[2]).merge(list_table.rows[3].cells[2])

            # row1 column3:
            ltr1c3 = list_table.rows[1].cells[3].add_paragraph()
            ltr1c3_run = ltr1c3.add_run()
            ltr1c3_run.bold = True
            ltr1c3_run.font.name = 'Liberation Sans'
            ltr1c3_run.font.size = Pt(10)
            list_table.rows[1].cells[3].merge(list_table.rows[1].cells[4])

            # row2 column0:
            ltr2c0 = list_table.rows[2].cells[0].add_paragraph()
            ltr2c0_run1_str = self.get_transl(cr, uid, "Currency:", partner.lang, context=context) + ' '
            ltr2c0_run1 = ltr2c0.add_run(ltr2c0_run1_str)
            ltr2c0_run1.bold = True
            ltr2c0_run1.font.name = 'Liberation Sans'
            ltr2c0_run1.font.size = Pt(10)
            ltr2c0_run2 = ltr2c0.add_run(company.currency_id.name)
            ltr2c0_run2.bold = True
            ltr2c0_run2.font.name = 'Liberation Sans'
            ltr2c0_run2.font.size = Pt(10)
            list_table.rows[2].cells[0].merge(list_table.rows[2].cells[1])

            # row2 column3:
            ltr2c3 = list_table.rows[2].cells[3].add_paragraph()
            ltr2c3_run_str = self.get_transl(cr, uid, "Currency:", partner.lang, context=context)
            ltr2c3_run = ltr2c3.add_run(ltr2c3_run_str)
            ltr2c3_run.bold = True
            ltr2c3_run.font.name = 'Liberation Sans'
            ltr2c3_run.font.size = Pt(10)
            list_table.rows[2].cells[3].merge(list_table.rows[2].cells[4])

            # row3 column0:
            ltr3c0 = list_table.rows[3].cells[0].add_paragraph()
            ltr3c0_run_str = self.get_transl(cr, uid, "Debit", partner.lang, context=context)
            ltr3c0_run = ltr3c0.add_run(ltr3c0_run_str)
            ltr3c0_run.bold = True
            ltr3c0_run.font.name = 'Liberation Sans'
            ltr3c0_run.font.size = Pt(12)

            # row3 column1:
            ltr3c1 = list_table.rows[3].cells[1].add_paragraph()
            ltr3c1_run_str = self.get_transl(cr, uid, "Credit", partner.lang, context=context)
            ltr3c1_run = ltr3c1.add_run(ltr3c1_run_str)
            ltr3c1_run.bold = True
            ltr3c1_run.font.name = 'Liberation Sans'
            ltr3c1_run.font.size = Pt(12)

            # row3 column3:
            ltr3c3 = list_table.rows[3].cells[3].add_paragraph()
            ltr3c3_run_str = self.get_transl(cr, uid, "Debit", partner.lang, context=context)
            ltr3c3_run = ltr3c3.add_run(ltr3c3_run_str)
            ltr3c3_run.bold = True
            ltr3c3_run.font.name = 'Liberation Sans'
            ltr3c3_run.font.size = Pt(12)

            # row3 column4:
            ltr3c4 = list_table.rows[3].cells[4].add_paragraph()
            ltr3c4_run_str = self.get_transl(cr, uid, "Credit", partner.lang, context=context)
            ltr3c4_run = ltr3c4.add_run(ltr3c4_run_str)
            ltr3c4_run.bold = True
            ltr3c4_run.font.name = 'Liberation Sans'
            ltr3c4_run.font.size = Pt(12)

            line_data = rabc(cr, uid, '', context).get_line_data(data['form']['date'], data['form']['type'], partner)

            # row4 column0:
            ltr4c0 = list_table.rows[4].cells[0].add_paragraph()
            ltr4c0_run_str = rml_parse(cr, uid, '', context).formatLang(line_data['total_debit'])
            ltr4c0_run = ltr4c0.add_run(ltr4c0_run_str)
            ltr4c0_run.bold = True
            ltr4c0_run.font.name = 'Liberation Sans'
            ltr4c0_run.font.size = Pt(11)

            # row4 column1:
            ltr4c1 = list_table.rows[4].cells[1].add_paragraph()
            ltr4c1_run_str = rml_parse(cr, uid, '', context).formatLang(line_data['total_credit'])
            ltr4c1_run = ltr4c1.add_run(ltr4c1_run_str)
            ltr4c1_run.bold = True
            ltr4c1_run.font.name = 'Liberation Sans'
            ltr4c1_run.font.size = Pt(11)

            # row4 column2:
            ltr4c2 = list_table.rows[4].cells[2].add_paragraph()
            ltr4c2_run = ltr4c2.add_run()
            ltr4c2_run.bold = True
            ltr4c2_run.font.name = 'Liberation Sans'
            ltr4c2_run.font.size = Pt(11)

            # row4 column3:
            ltr4c3 = list_table.rows[4].cells[3].add_paragraph()
            ltr4c3_run = ltr4c3.add_run()
            ltr4c3_run.bold = True
            ltr4c3_run.font.name = 'Liberation Sans'
            ltr4c3_run.font.size = Pt(11)

            # row4 column4:
            ltr4c4 = list_table.rows[4].cells[4].add_paragraph()
            ltr4c4_run = ltr4c4.add_run()
            ltr4c4_run.bold = True
            ltr4c4_run.font.name = 'Liberation Sans'
            ltr4c4_run.font.size = Pt(11)

            # row5 column0:
            ltr5c0_1 = list_table.rows[5].cells[0].add_paragraph()
            ltr5c0_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ltr5c0_1_run_str = self.get_transl(cr, uid, "To prevent disagreements the following additional journal entries need to be made", partner.lang, context=context)
            ltr5c0_1_run = ltr5c0_1.add_run(ltr5c0_1_run_str)
            ltr5c0_1_run.bold = True
            ltr5c0_1_run.font.name = 'Liberation Sans'
            ltr5c0_1_run.font.size = Pt(10)
            ltr5c0_2 = list_table.rows[5].cells[0].add_paragraph()
            ltr5c0_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ltr5c0_2_run_str = self.get_transl(cr, uid, "(declare disagreement and substantiation - document numbers and dates)", partner.lang, context=context)
            ltr5c0_2_run = ltr5c0_2.add_run(ltr5c0_2_run_str)
            ltr5c0_2_run.bold = True
            ltr5c0_2_run.font.name = 'Liberation Sans'
            ltr5c0_2_run.font.size = Pt(10)
            list_table.rows[5].cells[0].merge(list_table.rows[5].cells[1]).merge(list_table.rows[5].cells[2]).merge(list_table.rows[5].cells[3]).merge(list_table.rows[5].cells[4])

            # row6 column0:
            ltr6c0 = list_table.rows[6].cells[0].add_paragraph()
            ltr6c0_run_str = self.get_transl(cr, uid, "Debit", partner.lang, context=context)
            ltr6c0_run = ltr6c0.add_run(ltr6c0_run_str)
            ltr6c0_run.bold = True
            ltr6c0_run.font.name = 'Liberation Sans'
            ltr6c0_run.font.size = Pt(10)

            # row6 column1:
            ltr6c1 = list_table.rows[6].cells[1].add_paragraph()
            ltr6c1_run_str = self.get_transl(cr, uid, "Credit", partner.lang, context=context)
            ltr6c1_run = ltr6c1.add_run(ltr6c1_run_str)
            ltr6c1_run.bold = True
            ltr6c1_run.font.name = 'Liberation Sans'
            ltr6c1_run.font.size = Pt(10)

            # row6 column2:
            ltr6c2 = list_table.rows[6].cells[2].add_paragraph()
            ltr6c2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ltr6c2_run_str = self.get_transl(cr, uid, "Substantiation", partner.lang, context=context)
            ltr6c2_run = ltr6c2.add_run(ltr6c2_run_str)
            ltr6c2_run.bold = True
            ltr6c2_run.font.name = 'Liberation Sans'
            ltr6c2_run.font.size = Pt(10)

            # row6 column3:
            ltr6c3 = list_table.rows[6].cells[3].add_paragraph()
            ltr6c3_run_str = self.get_transl(cr, uid, "Debit", partner.lang, context=context)
            ltr6c3_run = ltr6c3.add_run(ltr6c3_run_str)
            ltr6c3_run.bold = True
            ltr6c3_run.font.name = 'Liberation Sans'
            ltr6c3_run.font.size = Pt(10)

            # row6 column4:
            ltr6c4 = list_table.rows[6].cells[4].add_paragraph()
            ltr6c4_run_str = self.get_transl(cr, uid, "Credit", partner.lang, context=context)
            ltr6c4_run = ltr6c4.add_run(ltr6c4_run_str)
            ltr6c4_run.bold = True
            ltr6c4_run.font.name = 'Liberation Sans'
            ltr6c4_run.font.size = Pt(10)

        file_data = StringIO.StringIO()
        document.save(file_data)
        file_data.seek(0)
        return file_data.read().decode('iso8859-4')

    def open_report(self, cr, uid, ids, context=None):
        data = {}
        data['ids'] = context.get('active_ids',[])
        data['model'] = 'res.partner'
        data['form'] = self.read(cr, uid, ids, ['date', 'type', 'accountant_id', 'format', 'file_name'], context=context)[0]
        for field in ['date', 'type', 'accountant_id', 'format', 'file_name']:
            if isinstance(data['form'][field], tuple):
                data['form'][field] = data['form'][field][0]
        used_context = self._build_contexts(cr, uid, ids, data, context=context)
        data['form']['used_context'] = used_context
        if data['form']['format'] == 'pdf':
            return {
                'type': 'ir.actions.report.xml',
                'report_name': 'l10n_lv_account.balance_comparison',
                'datas': data,
                'context': {
                    'active_ids': context.get('active_ids',[]),
                    'active_model': 'res.partner'
                }
            }
        if data['form']['format'] == 'docx':
            file_save = self.make_docx_data(cr, uid, context.get('active_ids',[]), data, context=context)
            file_save_data = base64.encodestring(file_save.encode('iso8859-4'))
            if data['form']['file_name']:
                file_name_list = data['form']['file_name'].split('.')
                format = file_name_list[-1]
                if format != 'docx':
                    if len(file_name_list) == 1:
                        file_name_list.append('docx')
                    else:
                        file_name_list[-1] = 'docx'
                file_name = '.'.join(file_name_list)
            else:
                file_name = 'Payment_Comparison.docx'
            self.write(cr, uid, ids[0], {
                'file_name': file_name,
                'file_save': file_save_data
            }, context=context)
            return {
                'name': _('Save document For Payment Comparison'),
                'res_id': ids[0],
                'context': context,
                'view_type': 'form',
                'view_mode': 'form',
                'res_model': 'account.balance.comparison.report',
                'views': [(False,'form')],
                'type': 'ir.actions.act_window',
                'target': 'new',
            }

# vim:expandtab:smartindent:tabstop=4:softtabstop=4:shiftwidth=4: